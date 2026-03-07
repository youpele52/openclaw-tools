# pyright: reportMissingImports=false
from pathlib import Path

import docx as _docx

from constants import DOCX_EXTENSIONS
from utils import ensure_output_dir, require_file, resolve_path


def read_docx(path_str: str) -> dict:
    """Extract text and basic metadata from a DOCX file."""
    path = resolve_path(path_str)
    err = require_file(path, DOCX_EXTENSIONS)
    if err:
        return {"error": err}

    try:
        doc = _docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        # Extract core properties
        props = doc.core_properties
        return {
            "file": str(path),
            "title": props.title or None,
            "author": props.author or None,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "paragraphs": len(paragraphs),
            "text": "\n".join(paragraphs),
        }
    except Exception as exc:
        return {"error": f"Could not read DOCX: {exc}"}


def write_docx(text: str, output_str: str) -> dict:
    """Write plain text to a DOCX file, one paragraph per line."""
    output = resolve_path(output_str)
    err = ensure_output_dir(output)
    if err:
        return {"error": err}

    try:
        doc = _docx.Document()
        lines = text.splitlines()
        para_count = 0
        for line in lines:
            doc.add_paragraph(line)
            para_count += 1

        doc.save(str(output))
        return {
            "output": str(output),
            "paragraphs": para_count,
            "chars": len(text),
        }
    except Exception as exc:
        return {"error": f"Could not write DOCX: {exc}"}
